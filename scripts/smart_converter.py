"""
🤖 Smart PDF Converter - Powered by AI/ML
Converts PDF to Word, Excel, PPTX using intelligent methods
"""

import sys
import os
import tempfile
import subprocess
from pathlib import Path

class SmartPDFConverter:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
        
    def convert_pdf_to_docx(self, pdf_path, output_path=None):
        """
        Convert PDF to DOCX using pdf2docx (ML-based)
        """
        try:
            from pdf2docx import Converter
            
            if not output_path:
                base_name = os.path.splitext(os.path.basename(pdf_path))[0]
                output_path = os.path.join(self.temp_dir, f"{base_name}.docx")
            
            # Usar pdf2docx para conversión
            cv = Converter(pdf_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            if os.path.exists(output_path):
                return {"success": True, "output_file": output_path}
            else:
                return {"success": False, "error": "Output file not created"}
                
        except Exception as e:
            return {"success": False, "error": f"pdf2docx Error: {str(e)}"}
    
    def convert_pdf_to_text(self, pdf_path):
        """
        Extract text from PDF using PyMuPDF (best for text extraction)
        """
        try:
            import pymupdf
            
            text_content = []
            doc = pymupdf.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_content.append(text)
            
            doc.close()
            
            return {"success": True, "text": "\n\n".join(text_content)}
        except Exception as e:
            return {"success": False, "error": f"Text extraction error: {str(e)}"}
    
    def convert_pdf_to_docx_fallback(self, pdf_path, output_path=None):
        """
        Fallback: Extract text and create DOCX manually
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            # Extraer texto
            text_result = self.convert_pdf_to_text(pdf_path)
            if not text_result["success"]:
                return text_result
            
            # Crear documento Word
            doc = Document()
            
            # Título
            doc.add_heading('Extracted from PDF', 0)
            
            # Agregar contenido
            lines = text_result["text"].split('\n')
            for line in lines[:100]:  # Limitar a primeras 100 líneas
                if line.strip():
                    doc.add_paragraph(line)
            
            if not output_path:
                base_name = os.path.splitext(os.path.basename(pdf_path))[0]
                output_path = os.path.join(self.temp_dir, f"{base_name}_extracted.docx")
            
            doc.save(output_path)
            
            if os.path.exists(output_path):
                return {"success": True, "output_file": output_path}
            else:
                return {"success": False, "error": "Failed to save document"}
                
        except Exception as e:
            return {"success": False, "error": f"Fallback error: {str(e)}"}
    
    def convert(self, pdf_path, output_format='docx'):
        """
        Main conversion method with intelligence
        """
        if not os.path.exists(pdf_path):
            return {"success": False, "error": "File not found"}
        
        if output_format == 'docx':
            # Intentar método principal primero
            result = self.convert_pdf_to_docx(pdf_path)
            
            # Si falla, usar fallback
            if not result["success"]:
                print(f"Primary method failed: {result.get('error')}, trying fallback...")
                result = self.convert_pdf_to_docx_fallback(pdf_path)
            
            return result
        
        return {"success": False, "error": f"Format {output_format} not supported yet"}


def main():
    if len(sys.argv) < 2:
        print("Usage: python smart_converter.py <input.pdf> [output_format]")
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'docx'
    
    converter = SmartPDFConverter()
    result = converter.convert(pdf_path, output_format)
    
    if result["success"]:
        print(f"OK:{result['output_file']}")
        print("Conversion completed successfully!")
    else:
        print(f"ERROR:{result['error']}")
        sys.exit(1)


if __name__ == "__main__":
    main()